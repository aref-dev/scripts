from docx import Document
from get_hours_from_specs import get_hours_from_specs
from format_datetime import format_datetime
from get_asset_details import get_asset_details

FINANCIAL_CUSTOM_PROPERTY_NAMES = ['Equipment_Holder',
                        'Purpose_of_Appraisal',
                        'Intended Use',
                        'Intended_User',
                        'Effective_Date',
                        'Report_Number',
                        'Market_Area',
                        'Attention',
                        'Report_Date',
                        'Reference',
                        'File_Number',
                        'Asset_1',
                        'Asset_1_VIN_Reported',
                        'Asset_1_VIN_Observed',
                        'Asset_1_VIN_Published',
                        'Asset_1_InspDate',
                        'Customer_Name',
                        'Customer_Address',
                        'Customer_Address_City']

class DocxReader:
    def __init__(self, file_path):
        self.properties = {}
        self.document = Document(file_path)
        self.set_custom_properties_vars()
        self.set_other_specs_vars()
        self.set_fmv_olv_flv_vars()
        self.extract_hours()

    def set_property(self, name, value):
        self.properties[name] = value
        
    def get_property(self, name):
        return self.properties.get(name, None)
    
    def set_custom_properties_vars(self):
        custom_properties = self.document.custom_properties
        custom_property_names = FINANCIAL_CUSTOM_PROPERTY_NAMES
        
        for prop_name in custom_property_names:
            if prop_name == 'Report_Number':
                prop_value = ''.join([char for char in custom_properties[prop_name] if char.isalpha()])
                self.set_property('Client', prop_value) 
            
            if prop_name in ('Effective_Date', 'Report_Date'):
                prop_value = format_datetime(custom_properties[prop_name], format_type='date')
                self.set_property(prop_name, prop_value)
                continue
            
            if prop_name == 'Asset_1':
                get_asset_details(custom_properties['Asset_1'])
                # year = None
                # make = None
                # model = None
                # description = None
                # self.set_property('Year', year)
                # self.set_property('Make', make)
                # self.set_property('Model', model)
                # self.set_property('Description', description)
                # continue
            
            prop_value = custom_properties[prop_name]
            self.set_property(prop_name, prop_value)

    def set_other_specs_vars(self):
        for paragraph in self.document.paragraphs: 
            for run in paragraph.runs: 
                if run.font.superscript: 
                    run.clear() 
                bullet_points = [] 
                collect = False

        for paragraph in self.document.paragraphs:
            if paragraph.style.name == 'Heading 1' and 'Asset Details' in paragraph.text:
                # Start collecting
                collect = True
            if collect and paragraph.style.name == 'List Paragraph':
                bullet_points.append(paragraph.text) 
            if paragraph.style.name == 'Heading 3' and 'Market Data ' in paragraph.text:
                break
        self.set_property('other_specs', bullet_points)
        
    def set_fmv_olv_flv_vars(self):
        headings = ['Asset Description (Serial Number)', 'Effective Date of Value', 'FMV', 'OLV', 'FLV']
        table_values = self.get_table(headings)[0]
        
        fmv_range = table_values[2]
        olv_range = table_values[3]
        flv_range = table_values[4]

        fmv_low, fmv_high = self.clean_and_split(fmv_range)
        olv_low, olv_high = self.clean_and_split(olv_range)
        flv_lo, flv_high = self.clean_and_split(flv_range)
        
        self.set_property('FMV - Low', fmv_low)
        self.set_property('FMV - High', fmv_high)
        self.set_property('OLV - Low', olv_low)
        self.set_property('OLV - High', olv_high)
        self.set_property('FLV - Low', flv_lo)
        self.set_property('FLV - High', flv_high)
            
    def extract_hours(self):
        hours = get_hours_from_specs(self.get_property('other_specs'))
        self.set_property('Hours', hours)
    
    def get_table(self, headings):
        for table in self.document.tables:
        # Check if the first row contains the specified headings
            first_row = [cell.text.strip() for cell in table.rows[0].cells]
            if first_row == headings:
                table_values = []
                for row in table.rows[1:]:  # Skip the header row
                    row_values = [cell.text.strip() for cell in row.cells]
                    table_values.append(row_values)
                return table_values
        
    def clean_and_split(self, value):
        cleaned_value = value.replace("$", "").replace(",", "")
        low, high = cleaned_value.split(" - ")
        return int(float(low)), int(float(high))
        
